import os
import json
import csv
import asyncio
from openai import AsyncOpenAI
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
client = AsyncOpenAI(api_key=api_key)

# Configuración
projects = [
    {
        "name": "TADL_RAG_SECRET_PROJECT",
        "description": "Proyecto Secreto TADL RAG",
        "status": "activo",
        "outlook": "¡El proyecto va excelente! ¡Estamos logrando avances increíbles!",
    },
    {
        "name": "UX_OPENAI_REFACTOR",
        "description": "Refactorización UX OpenAI",
        "status": "activo",
        "outlook": "El proyecto va bien. Nada destacable.",
    },
    {
        "name": "ITBA_NEW_AUTHENTICATION_SERVER",
        "description": "Nuevo Servidor de Autenticación ITBA",
        "status": "activo",
        "outlook": "El proyecto muestra muy poco progreso. Muchas cosas no están funcionando como se esperaba.",
    },
]
num_files_per_project = 5
output_dir = "backend/uploads"


async def generate_ai_content(prompt: str, model: str = "gpt-4.1") -> str:
    """Llama a OpenAI ChatCompletion para generar contenido asincrónicamente."""
    system_message = {
        "role": "system",
        "content": "Eres un generador de datos sintéticos para documentación de proyectos de software. Genera contenido realista y detallado que incluya: actividades recientes, problemas/bloqueos, interacciones con otros equipos, KPIs específicos y tareas planificadas. Usa lenguaje técnico apropiado. Sin formato extra, solo contenido.",
    }
    response = await client.responses.create(
        model=model,
        input=[system_message, {"role": "user", "content": prompt}],
        temperature=0.7,
    )
    return response.output_text.strip()


async def create_pdf(project: str, i: int, timestamp: str, content: str) -> str:
    """Crea un archivo PDF con el contenido dado."""
    project_dir = os.path.join(output_dir, project)
    os.makedirs(project_dir, exist_ok=True)

    pdf_filename = f"{project}_summary_{i}_{timestamp}.pdf"
    full_path = os.path.join(project_dir, pdf_filename)

    c = canvas.Canvas(full_path, pagesize=letter)
    text_obj = c.beginText(40, 720)
    for line in content.split("\n"):
        text_obj.textLine(line)
    c.drawText(text_obj)
    c.save()
    return pdf_filename


async def create_docx(project: str, i: int, timestamp: str, content: str) -> str:
    """Crea un archivo DOCX con el contenido dado."""
    project_dir = os.path.join(output_dir, project)
    os.makedirs(project_dir, exist_ok=True)

    docx_filename = f"{project}_brief_{i}_{timestamp}.docx"
    full_path = os.path.join(project_dir, docx_filename)

    def _create_docx():
        doc = Document()
        doc.add_heading(f"{project} Brief #{i}", level=1)
        for para in content.split("\n\n"):
            # Check if paragraph contains markdown headings or lists
            if (
                para.strip().startswith("#")
                or para.strip().startswith("*")
                or para.strip().startswith("-")
            ):
                # Add heading or list item directly
                doc.add_paragraph(para)
            else:
                # Regular paragraph
                doc.add_paragraph(para)
        doc.save(full_path)

    await asyncio.to_thread(_create_docx)
    return docx_filename


async def create_json(project: str, i: int, timestamp: str, content: str) -> str:
    """Crea un archivo JSON con una conversación estilo Slack."""
    project_dir = os.path.join(output_dir, project)
    os.makedirs(project_dir, exist_ok=True)

    json_filename = f"{project}_chat_{i}_{timestamp}.json"
    full_path = os.path.join(project_dir, json_filename)

    messages = json.loads(content)

    await asyncio.to_thread(
        lambda: json.dump(messages, open(full_path, "w", encoding="utf-8"), indent=2)
    )
    return json_filename


async def create_csv(project: str, i: int, timestamp: str, content: str) -> str:
    """Crea un archivo CSV con el contenido dado."""
    project_dir = os.path.join(output_dir, project)
    os.makedirs(project_dir, exist_ok=True)

    csv_filename = f"{project}_metrics_{i}_{timestamp}.csv"
    full_path = os.path.join(project_dir, csv_filename)

    await asyncio.to_thread(
        lambda: open(full_path, "w", newline="", encoding="utf-8").write(content)
    )
    return csv_filename


async def process_file_set(project: dict, i: int):
    """Procesa un conjunto completo de archivos para un proyecto."""
    project_name = project["name"]
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

    context = f"Contexto del proyecto: '{project['description']}'. Estado: '{project['status']}'. Perspectiva: '{project['outlook']}'"

    pdf_content_task = generate_ai_content(
        f"Escribe un resumen detallado (250-300 palabras) del proyecto '{project_name}'. {context}. Incluye: estado actual, logros recientes, problemas identificados y próximos pasos. Menciona KPIs específicos, avances en el sprint actual y cualquier interacción con otros equipos. Usa términos técnicos y fechas concretas."
    )

    docx_content_task = generate_ai_content(
        f"Crea un Brief completo para '{project_name}'. {context}. Utiliza las siguientes secciones en formato markdown:\n# Resumen Ejecutivo\n## Objetivos del Sprint\n## Actividades Completadas\n## Problemas Técnicos\n## Interacciones con Otros Equipos\n## Métricas y KPIs\n## Plan de Trabajo\nIncluye detalles técnicos específicos, menciona tecnologías relevantes, y asigna tareas a miembros del equipo."
    )

    json_content_task = generate_ai_content(
        f"Genera un arreglo JSON de 8-10 mensajes realistas entre miembros del equipo discutiendo '{project_name}'. {context}. Incluye discusiones sobre problemas técnicos, logros, bloqueos y próximos pasos. Cada mensaje debe tener `user` (nombres reales), `text` (con menciones, emojis y referencias técnicas), `ts` (timestamp) y `reactions` (opcional). Representa una conversación técnica auténtica."
    )

    csv_content_task = generate_ai_content(
        f"Proporciona métricas detalladas del proyecto '{project_name}'. {context}. Formato CSV con las columnas: fecha,categoría,métrica,valor,tendencia,responsable. Incluye 8-10 filas con métricas como velocidad de sprint, bugs por release, cobertura de código, tiempo de ciclo, etc. Usa valores realistas y fechas recientes."
    )

    pdf_content, docx_content, json_content, csv_content = await asyncio.gather(
        pdf_content_task, docx_content_task, json_content_task, csv_content_task
    )

    pdf_task = create_pdf(project_name, i, timestamp, pdf_content)
    docx_task = create_docx(project_name, i, timestamp, docx_content)
    json_task = create_json(project_name, i, timestamp, json_content)
    csv_task = create_csv(project_name, i, timestamp, csv_content)

    pdf_filename, docx_filename, json_filename, csv_filename = await asyncio.gather(
        pdf_task, docx_task, json_task, csv_task
    )

    project_dir = os.path.join(output_dir, project_name)
    print(
        f"Archivos generados para {project_name} en {project_dir}: {pdf_filename}, {docx_filename}, {json_filename}, {csv_filename}"
    )


async def main():
    """Punto de entrada principal del script."""
    os.makedirs(output_dir, exist_ok=True)
    print(
        f"Guardando todos los archivos generados en el directorio: {os.path.abspath(output_dir)}"
    )

    tasks = []
    for project in projects:
        for i in range(1, num_files_per_project + 1):
            tasks.append(process_file_set(project, i))

    await asyncio.gather(*tasks)
    print(
        f"Completada la generación de {len(tasks)} conjuntos de archivos en {output_dir}"
    )


if __name__ == "__main__":
    asyncio.run(main())
